"""
File Processor Service
Handles file upload, validation, and text extraction from PDF and DOCX files.
"""

import os
import logging
from typing import Tuple
from werkzeug.datastructures import FileStorage
from werkzeug.utils import secure_filename
import PyPDF2
import pdfplumber
import docx
from pdfminer.high_level import extract_text as pdfminer_extract

from utils.exceptions import FileProcessingError, FileValidationError
from utils.validators import sanitize_filename, is_text_extractable

# Configure logging
logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)


class FileProcessor:
    """
    Processes uploaded resume files and extracts text content.

    Supports:
    - PDF files (using multiple libraries with fallback)
    - DOCX files (using python-docx)
    """

    def __init__(self, upload_folder: str):
        """
        Initialize file processor

        Args:
            upload_folder: Directory to store uploaded files temporarily
        """
        self.upload_folder = upload_folder
        os.makedirs(upload_folder, exist_ok=True)

    def process_file(self, file: FileStorage) -> Tuple[str, dict]:
        """
        Process uploaded file and extract text

        Args:
            file: Uploaded file object

        Returns:
            Tuple of (extracted_text, metadata)

        Raises:
            FileProcessingError: If file processing fails
        """
        # Save file temporarily
        filepath = self._save_file(file)

        try:
            # Determine file type and extract text
            extension = self._get_extension(file.filename)

            if extension == 'pdf':
                text = self._extract_text_pdf(filepath)
            elif extension == 'docx':
                text = self._extract_text_docx(filepath)
            else:
                raise FileProcessingError(
                    f"Unsupported file type: {extension}",
                    {'extension': extension}
                )

            # Validate extracted text
            if not is_text_extractable(text):
                raise FileProcessingError(
                    "Failed to extract meaningful text from the file. The file may be corrupted, password-protected, or contain only images.",
                    {'filename': file.filename}
                )

            # Generate metadata
            metadata = {
                'filename': file.filename,
                'size_bytes': os.path.getsize(filepath),
                'size_kb': round(os.path.getsize(filepath) / 1024, 2),
                'extension': extension,
                'word_count': len(text.split()),
                'char_count': len(text)
            }

            logger.info(f"Successfully processed file: {file.filename} ({metadata['word_count']} words)")

            return text, metadata

        finally:
            # Clean up temporary file
            self._cleanup_file(filepath)

    def _save_file(self, file: FileStorage) -> str:
        """
        Save uploaded file temporarily

        Args:
            file: Uploaded file object

        Returns:
            Path to saved file
        """
        # Sanitize and secure filename
        filename = sanitize_filename(file.filename)
        filename = secure_filename(filename)

        # Add timestamp to avoid collisions
        import time
        timestamp = int(time.time() * 1000)
        name, ext = os.path.splitext(filename)
        filename = f"{name}_{timestamp}{ext}"

        filepath = os.path.join(self.upload_folder, filename)

        try:
            file.save(filepath)
            logger.info(f"Saved file: {filepath}")
            return filepath
        except Exception as e:
            logger.error(f"Failed to save file: {str(e)}")
            raise FileProcessingError(
                f"Failed to save uploaded file: {str(e)}",
                {'filename': file.filename}
            )

    def _extract_text_pdf(self, filepath: str) -> str:
        """
        Extract text from PDF file using multiple methods with fallback

        Args:
            filepath: Path to PDF file

        Returns:
            Extracted text

        Raises:
            FileProcessingError: If all extraction methods fail
        """
        text = ""

        # Method 1: Try pdfplumber (best for complex layouts)
        try:
            text = self._extract_with_pdfplumber(filepath)
            if is_text_extractable(text):
                logger.info(f"Extracted text using pdfplumber: {len(text)} characters")
                return text
        except Exception as e:
            logger.warning(f"pdfplumber extraction failed: {str(e)}")

        # Method 2: Try PyPDF2 (good for simple PDFs)
        try:
            text = self._extract_with_pypdf2(filepath)
            if is_text_extractable(text):
                logger.info(f"Extracted text using PyPDF2: {len(text)} characters")
                return text
        except Exception as e:
            logger.warning(f"PyPDF2 extraction failed: {str(e)}")

        # Method 3: Try pdfminer (most robust, slower)
        try:
            text = self._extract_with_pdfminer(filepath)
            if is_text_extractable(text):
                logger.info(f"Extracted text using pdfminer: {len(text)} characters")
                return text
        except Exception as e:
            logger.warning(f"pdfminer extraction failed: {str(e)}")

        # All methods failed
        raise FileProcessingError(
            "Failed to extract text from PDF. The file may be scanned, password-protected, or corrupted.",
            {'filepath': filepath}
        )

    def _extract_with_pdfplumber(self, filepath: str) -> str:
        """Extract text using pdfplumber"""
        text_parts = []
        with pdfplumber.open(filepath) as pdf:
            for page in pdf.pages:
                page_text = page.extract_text()
                if page_text:
                    text_parts.append(page_text)
        return "\n\n".join(text_parts)

    def _extract_with_pypdf2(self, filepath: str) -> str:
        """Extract text using PyPDF2"""
        text_parts = []
        with open(filepath, 'rb') as file:
            pdf_reader = PyPDF2.PdfReader(file)
            for page in pdf_reader.pages:
                page_text = page.extract_text()
                if page_text:
                    text_parts.append(page_text)
        return "\n\n".join(text_parts)

    def _extract_with_pdfminer(self, filepath: str) -> str:
        """Extract text using pdfminer"""
        return pdfminer_extract(filepath)

    def _extract_text_docx(self, filepath: str) -> str:
        """
        Extract text from DOCX file

        Args:
            filepath: Path to DOCX file

        Returns:
            Extracted text

        Raises:
            FileProcessingError: If extraction fails
        """
        try:
            doc = docx.Document(filepath)

            # Extract text from paragraphs
            text_parts = []
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    text_parts.append(paragraph.text)

            # Extract text from tables
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        if cell.text.strip():
                            text_parts.append(cell.text)

            text = "\n".join(text_parts)

            if not is_text_extractable(text):
                raise FileProcessingError(
                    "Failed to extract meaningful text from DOCX file",
                    {'filepath': filepath}
                )

            logger.info(f"Extracted text from DOCX: {len(text)} characters")
            return text

        except Exception as e:
            logger.error(f"DOCX extraction failed: {str(e)}")
            raise FileProcessingError(
                f"Failed to extract text from DOCX file: {str(e)}",
                {'filepath': filepath}
            )

    def _cleanup_file(self, filepath: str) -> None:
        """
        Delete temporary file

        Args:
            filepath: Path to file to delete
        """
        try:
            if os.path.exists(filepath):
                os.remove(filepath)
                logger.info(f"Cleaned up file: {filepath}")
        except Exception as e:
            logger.warning(f"Failed to clean up file {filepath}: {str(e)}")

    @staticmethod
    def _get_extension(filename: str) -> str:
        """Get file extension"""
        if '.' not in filename:
            return ''
        return filename.rsplit('.', 1)[1].lower()


def clean_text(text: str) -> str:
    """
    Clean and normalize extracted text

    Args:
        text: Raw extracted text

    Returns:
        Cleaned text
    """
    # Remove excessive whitespace
    lines = [line.strip() for line in text.split('\n')]
    lines = [line for line in lines if line]

    # Join lines
    text = "\n".join(lines)

    # Normalize unicode characters
    text = text.replace('\u2019', "'")  # Right single quotation mark
    text = text.replace('\u2018', "'")  # Left single quotation mark
    text = text.replace('\u201c', '"')  # Left double quotation mark
    text = text.replace('\u201d', '"')  # Right double quotation mark
    text = text.replace('\u2013', '-')  # En dash
    text = text.replace('\u2014', '-')  # Em dash
    text = text.replace('\u2022', '*')  # Bullet point
    text = text.replace('\xa0', ' ')    # Non-breaking space

    return text.strip()
